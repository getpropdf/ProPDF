"""ProPDF Engine test suite — runs every pipeline against the corpus and
scores conversion quality 0–100. Outputs are actually opened and validated;
totals are checked against ground truth to the paisa."""
import json
import os
import sys
import time
import traceback
from decimal import Decimal

sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), ".."))

CORPUS = os.environ.get("PROPDF_CORPUS_DIR",
                        os.path.join(os.path.dirname(os.path.abspath(__file__)), "corpus"))
OUT = os.environ.get("PROPDF_TEST_OUT", "/tmp/propdf_test_out")
os.makedirs(OUT, exist_ok=True)

RESULTS = []


def record(name, score, detail):
    RESULTS.append({"test": name, "score": score, "detail": detail})
    print(f"  [{score:5.1f}] {name} — {detail}")


def d(x):
    return Decimal(str(x))


def score_statement(name, res, gt):
    s = 0.0
    notes = []
    # completeness 30
    if res.get("rows") == gt["n_rows"]:
        s += 30
    else:
        s += max(0, 30 - 3 * abs(res.get("rows", 0) - gt["n_rows"]))
        notes.append(f"rows {res.get('rows')}/{gt['n_rows']}")
    # numeric accuracy 50
    ok = 0
    for k_res, k_gt in [("total_debits", "total_debits"), ("total_credits", "total_credits"),
                        ("closing_balance", "closing"), ("opening_balance", "opening")]:
        try:
            if res.get(k_res) is not None and d(res[k_res]) == d(gt[k_gt]):
                ok += 1
            else:
                notes.append(f"{k_res}={res.get(k_res)} gt={gt[k_gt]}")
        except Exception:
            notes.append(f"{k_res} unparseable")
    s += 12.5 * ok
    # continuity 20
    s += 20 * (res.get("continuity_confidence", 0) / 100.0)
    return min(100.0, round(s, 1)), ("; ".join(notes) or "exact")


def test_excel():
    from core.excel import convert_pdf_to_excel
    from openpyxl import load_workbook

    print("\n== PDF → Excel ==")
    for f in ["bank_statement_ruled_multipage", "bank_statement_borderless",
              "bank_statement_scanned"]:
        try:
            t0 = time.time()
            res = convert_pdf_to_excel(os.path.join(CORPUS, f + ".pdf"),
                                       os.path.join(OUT, f + ".xlsx"))
            gt_file = f if f != "bank_statement_scanned" else "bank_statement_borderless"
            gt = json.load(open(os.path.join(CORPUS, gt_file + ".pdf.json")))
            load_workbook(os.path.join(OUT, f + ".xlsx"))  # integrity check
            sc, note = score_statement(f, res, gt)
            record(f"pdf2excel/{f}", sc, f"{note} ({time.time()-t0:.1f}s)")
        except Exception as e:
            traceback.print_exc()
            record(f"pdf2excel/{f}", 0, f"EXCEPTION {e}")

    # trial balance: generic mode numeric check
    try:
        res = convert_pdf_to_excel(os.path.join(CORPUS, "trial_balance.pdf"),
                                   os.path.join(OUT, "trial_balance.xlsx"))
        gt = json.load(open(os.path.join(CORPUS, "trial_balance.pdf.json")))
        ws = load_workbook(os.path.join(OUT, "trial_balance.xlsx")).active
        rows = [[c.value for c in r] for r in ws.iter_rows()]
        tot = next((r for r in rows if r and str(r[0]).strip().upper() == "TOTAL"), None)
        okd = tot and any(v is not None and d(v) == d(gt["total_debit"]) for v in tot[1:])
        okc = tot and any(v is not None and d(v) == d(gt["total_credit"]) for v in tot[1:])
        sc = 40 + (30 if okd else 0) + (30 if okc else 0)
        record("pdf2excel/trial_balance", sc,
               f"type={res['doc_type']} totals {'exact' if okd and okc else 'MISMATCH'}")
    except Exception as e:
        record("pdf2excel/trial_balance", 0, f"EXCEPTION {e}")

    # rotated scan
    try:
        res = convert_pdf_to_excel(os.path.join(CORPUS, "trial_balance_rotated.pdf"),
                                   os.path.join(OUT, "tb_rot.xlsx"))
        sc = 0
        if res["rows"] >= 12:
            sc += 60
        conf = res.get("mean_ocr_confidence") or 0
        sc += min(40, conf * 0.4)
        record("pdf2excel/trial_balance_rotated", round(sc, 1),
               f"rows={res['rows']} ocr_conf={conf}")
    except Exception as e:
        record("pdf2excel/trial_balance_rotated", 0, f"EXCEPTION {e}")


def test_word():
    from core.word_out import convert_pdf_to_word
    from docx import Document

    print("\n== PDF → Word ==")
    try:
        convert_pdf_to_word(os.path.join(CORPUS, "text_report.pdf"),
                            os.path.join(OUT, "report_layout.docx"), mode="layout")
        doc = Document(os.path.join(OUT, "report_layout.docx"))
        txt = "\n".join(p.text for p in doc.paragraphs)
        sc = 0
        if "Annual Compliance Review" in txt:
            sc += 50
        if "section 44AA" in txt:
            sc += 30
        if len(txt) > 2000:
            sc += 20
        record("pdf2word/layout_text_report", sc, f"{len(txt)} chars extracted")
    except Exception as e:
        record("pdf2word/layout_text_report", 0, f"EXCEPTION {e}")

    try:
        r = convert_pdf_to_word(os.path.join(CORPUS, "invoice_scanned_lowq.pdf"),
                                os.path.join(OUT, "invoice_edit.docx"), mode="editable")
        doc = Document(os.path.join(OUT, "invoice_edit.docx"))
        txt = " ".join(p.text for p in doc.paragraphs)
        sc = 0
        if "TAX INVOICE" in txt:
            sc += 40
        if "27ABCDE1234F" in txt.replace(" ", ""):
            sc += 30
        conf = r.get("mean_ocr_confidence") or 0
        sc += min(30, conf * 0.3)
        record("pdf2word/editable_scanned_invoice", round(sc, 1),
               f"auto-OCR conf={conf}")
    except Exception as e:
        record("pdf2word/editable_scanned_invoice", 0, f"EXCEPTION {e}")


def test_office():
    from core.office import convert_office_to_pdf
    import fitz

    print("\n== Office → PDF ==")
    from docx import Document as Doc
    p = os.path.join(OUT, "letter.docx")
    dd = Doc()
    dd.add_heading("Engagement Letter", 0)
    dd.add_paragraph("Audit engagement for FY 2025-26. Fee: Rs. 1,50,000.")
    dd.save(p)
    try:
        convert_office_to_pdf(p, os.path.join(OUT, "letter.pdf"))
        t = fitz.open(os.path.join(OUT, "letter.pdf"))[0].get_text()
        sc = 100 if "Engagement Letter" in t and "1,50,000" in t else 40
        record("office2pdf/docx", sc, "text verified in output PDF")
    except Exception as e:
        record("office2pdf/docx", 0, f"EXCEPTION {e}")


def test_ops():
    from core import pdfops as P
    import fitz

    print("\n== PDF ops ==")
    src = os.path.join(CORPUS, "bank_statement_ruled_multipage.pdf")

    try:
        P.protect(src, os.path.join(OUT, "prot.pdf"), user_pw="secret1")
        try:
            fitz.open(os.path.join(OUT, "prot.pdf"))[0].get_text()
            record("protect/aes256", 0, "opened without password!")
        except Exception:
            P.unlock(os.path.join(OUT, "prot.pdf"), os.path.join(OUT, "unlock.pdf"),
                     password="secret1")
            ok = len(fitz.open(os.path.join(OUT, "unlock.pdf"))[0].get_text()) > 100
            record("protect+unlock/aes256", 100 if ok else 40, "round-trip verified")
    except Exception as e:
        record("protect+unlock/aes256", 0, f"EXCEPTION {e}")

    try:
        data = open(src, "rb").read()
        dmg = os.path.join(OUT, "damaged.pdf")
        open(dmg, "wb").write(data[:-350])
        P.repair(dmg, os.path.join(OUT, "repaired.pdf"))
        n = fitz.open(os.path.join(OUT, "repaired.pdf")).page_count
        record("repair/truncated_xref", 100 if n >= 4 else 40, f"{n} pages recovered")
    except Exception as e:
        record("repair/truncated_xref", 0, f"EXCEPTION {e}")

    try:
        r = P.compress(os.path.join(CORPUS, "bank_statement_scanned.pdf"),
                       os.path.join(OUT, "comp.pdf"), "ebook")
        ok = fitz.open(os.path.join(OUT, "comp.pdf")).page_count == 2
        sc = 100 if ok and r["saved_pct"] > 30 else (60 if ok else 0)
        record("compress/scanned", sc, f"{r['saved_pct']}% saved, {r['engine']}")
    except Exception as e:
        record("compress/scanned", 0, f"EXCEPTION {e}")

    try:
        r = P.deskew_pdf(os.path.join(CORPUS, "trial_balance_rotated.pdf"),
                         os.path.join(OUT, "deskewed.pdf"))
        record("deskew/rotated90", 100 if r["corrected"] >= 1 else 30,
               f"corrected {r['corrected']}/{r['pages']}")
    except Exception as e:
        record("deskew/rotated90", 0, f"EXCEPTION {e}")

    try:
        P.strip_metadata(src, os.path.join(OUT, "nometa.pdf"))
        md = fitz.open(os.path.join(OUT, "nometa.pdf")).metadata or {}
        clean = not (md.get("author") or md.get("producer") or md.get("creator"))
        record("strip_metadata", 100 if clean else 50, str({k: v for k, v in md.items() if v}))
    except Exception as e:
        record("strip_metadata", 0, f"EXCEPTION {e}")


def test_ocr_searchable():
    from core.ocr import make_searchable_pdf
    import fitz

    print("\n== Searchable OCR PDF ==")
    try:
        r = make_searchable_pdf(os.path.join(CORPUS, "bank_statement_scanned.pdf"),
                                os.path.join(OUT, "searchable.pdf"))
        t = fitz.open(os.path.join(OUT, "searchable.pdf"))[0].get_text().upper()
        sc = 0
        if "OPENING BALANCE" in t:
            sc += 60
        conf = r.get("mean_confidence") or 0
        sc += min(40, conf * 0.4)
        record("ocr/searchable_pdf", round(sc, 1), f"conf={conf}")
    except Exception as e:
        record("ocr/searchable_pdf", 0, f"EXCEPTION {e}")


def test_bigfile():
    """100+ page performance check."""
    import fitz
    from core.excel import convert_pdf_to_excel

    print("\n== 120-page stress test ==")
    big = os.path.join(OUT, "big.pdf")
    src = fitz.open(os.path.join(CORPUS, "bank_statement_ruled_multipage.pdf"))
    doc = fitz.open()
    while doc.page_count < 120:
        doc.insert_pdf(src)
    doc.save(big)
    doc.close()
    t0 = time.time()
    try:
        res = convert_pdf_to_excel(big, os.path.join(OUT, "big.xlsx"))
        dt = time.time() - t0
        sc = 100 if res["rows"] > 2000 and dt < 300 else 60
        record("stress/120_pages", sc, f"{res['rows']} rows in {dt:.1f}s")
    except Exception as e:
        record("stress/120_pages", 0, f"EXCEPTION {e}")


if __name__ == "__main__":
    only = sys.argv[1] if len(sys.argv) > 1 else "all"
    if only in ("all", "excel"):
        test_excel()
    if only in ("all", "word"):
        test_word()
    if only in ("all", "office"):
        test_office()
    if only in ("all", "ops"):
        test_ops()
    if only in ("all", "ocr"):
        test_ocr_searchable()
    if only in ("all", "big"):
        test_bigfile()
    print("\n==== SUMMARY ====")
    avg = sum(r["score"] for r in RESULTS) / max(1, len(RESULTS))
    for r in RESULTS:
        print(f"  {r['score']:5.1f}  {r['test']}")
    print(f"  MEAN QUALITY SCORE: {avg:.1f}/100 over {len(RESULTS)} tests")
    json.dump(RESULTS, open(os.path.join(OUT, "results.json"), "w"), indent=1)
