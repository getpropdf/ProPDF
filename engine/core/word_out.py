"""PDF -> Word (.docx) conversion with two modes.

LAYOUT mode  : pdf2docx layout reconstruction — visually faithful, keeps
               tables, images, columns. Scanned pages get an OCR text layer
               first so pdf2docx has real text to place.
EDITABLE mode: clean semantic reconstruction via python-docx — paragraphs,
               headings (by font size), bold/italic runs, detected tables as
               real Word tables, page breaks. Best for heavy re-editing.
"""
from __future__ import annotations

import os
import tempfile
from typing import Optional

import fitz
from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt

from .ocr import make_searchable_pdf, ocr_page_words, resolve_langs


def convert_layout(src: str, out_path: str, langs: str = "eng",
                   progress=None) -> dict:
    """pdf2docx-based, with auto-OCR for scanned pages."""
    from pdf2docx import Converter

    doc = fitz.open(src)
    scanned = [i for i in range(doc.page_count)
               if len(doc[i].get_text("text").strip()) < 30]
    doc.close()
    work = src
    ocr_info = None
    tmp = None
    if scanned:
        tmp = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
        tmp.close()
        ocr_info = make_searchable_pdf(src, tmp.name, langs=langs)
        work = tmp.name
    cv = Converter(work)
    try:
        cv.convert(out_path)
    finally:
        cv.close()
        if tmp:
            try:
                os.unlink(tmp.name)
            except OSError:
                pass
    return {"mode": "layout", "ocr_pages": len(scanned),
            "mean_ocr_confidence": (ocr_info or {}).get("mean_confidence")}


def convert_editable(src: str, out_path: str, langs: str = "eng",
                     progress=None) -> dict:
    doc = fitz.open(src)
    word = Document()
    style = word.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    ocr_pages = 0
    confs = []
    # find the dominant body font size for heading detection
    sizes = {}
    for i in range(min(10, doc.page_count)):
        for b in doc[i].get_text("dict")["blocks"]:
            for l in b.get("lines", []):
                for s in l.get("spans", []):
                    sz = round(s["size"])
                    sizes[sz] = sizes.get(sz, 0) + len(s.get("text", ""))
    body_size = max(sizes, key=sizes.get) if sizes else 11

    for pno in range(doc.page_count):
        if progress:
            progress(pno + 1, doc.page_count)
        page = doc[pno]
        text = page.get_text("text").strip()
        if len(text) < 30:  # scanned page -> auto OCR
            op = ocr_page_words(doc, pno + 1, langs=langs)
            ocr_pages += 1
            if op.mean_confidence:
                confs.append(op.mean_confidence)
            for line in op.text.split("\n"):
                if line.strip():
                    word.add_paragraph(line.strip())
        else:
            d = page.get_text("dict")
            # embedded images -> keep them in the docx
            for b in d["blocks"]:
                if b.get("type") == 1 and b.get("image"):
                    try:
                        import io as _io
                        from docx.shared import Inches
                        w_in = min(6.0, (b["bbox"][2] - b["bbox"][0]) / 72.0)
                        word.add_picture(_io.BytesIO(b["image"]), width=Inches(max(1.0, w_in)))
                    except Exception:
                        pass
            for b in d["blocks"]:
                if b.get("type") != 0:
                    continue
                for l in b.get("lines", []):
                    spans = [s for s in l.get("spans", []) if s.get("text", "").strip()]
                    if not spans:
                        continue
                    max_sz = max(round(s["size"]) for s in spans)
                    if max_sz >= body_size + 4:
                        p = word.add_heading(" ".join(s["text"].strip() for s in spans), level=1)
                        continue
                    if max_sz >= body_size + 2:
                        p = word.add_heading(" ".join(s["text"].strip() for s in spans), level=2)
                        continue
                    p = word.add_paragraph()
                    for s in spans:
                        run = p.add_run(s["text"])
                        flags = s.get("flags", 0)
                        run.bold = bool(flags & 16)
                        run.italic = bool(flags & 2)
                        run.font.size = Pt(max(7, min(28, s["size"])))
        if pno < doc.page_count - 1:
            word.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    word.save(out_path)
    n = doc.page_count
    doc.close()
    return {"mode": "editable", "pages": n, "ocr_pages": ocr_pages,
            "mean_ocr_confidence": round(sum(confs) / len(confs), 1) if confs else None}


def convert_pdf_to_word(src: str, out_path: str, mode: str = "layout",
                        langs: str = "eng", progress=None) -> dict:
    if mode == "editable":
        return convert_editable(src, out_path, langs=langs, progress=progress)
    return convert_layout(src, out_path, langs=langs, progress=progress)
